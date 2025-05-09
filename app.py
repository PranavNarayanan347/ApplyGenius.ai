from flask import Flask, render_template, request
from docx import Document
from flask import send_file
from dotenv import load_dotenv
import os 
import pdfplumber
import docx
import openai
from flask import after_this_request
import requests
from bs4 import BeautifulSoup
import json
from flask import redirect
from datetime import datetime

load_dotenv()

openai.api_key = os.getenv("OPENAI_API_KEY")
app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

@app.route('/')
def home():
    return render_template('landing.html', year=datetime.now().year)

def extract_text_from_pdf(path):
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(path):
    doc = docx.Document(path)
    return "\n".join([p.text for p in doc.paragraphs])

from openai import OpenAI

client = OpenAI()

def rewrite_resume(resume_text, job_description):
    prompt = f"""
You are a resume optimization expert.

Your job is to rewrite only the content of the resume to better match the job description using relevant keywords and phrasing. However:

- ✅ Do NOT remove or change any section headings.
- ✅ Do NOT change the structure of the resume.
- ✅ Do NOT remove any bullet points unless necessary.
- ✅ Keep the same order of content (e.g., Experience, Education, Skills).
- ✅ Only change wording to better align with the job description.

--- Resume ---
{resume_text}

--- Job Description ---
{job_description}
"""

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",  # or "gpt-3.5-turbo"
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1000
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"❌ GPT failed: {str(e)}"

def scrape_webpage(url):
    Source = requests.get(url)
    soup = BeautifulSoup(Source.text, 'html.parser')
    for script in soup(["script", "style"]):  # remove scripts and styles
        script.decompose()
    page_text = soup.get_text()
    return page_text

def extract_job_description_from_text(page_text):
    prompt = f"""
You are a job listing parser.

Below is the full text content from a job listing webpage, including irrelevant text like navigation, footer, and metadata.

Your task:
- Extract the company name, job title, and job description
- Focus on the main content of the job description
- Extract only the job description section
- Include responsibilities, qualifications, and company-provided details
- Ignore navigation links, unrelated paragraphs, or footers

Here is the webpage text:
\"\"\"
{page_text}
\"\"\"

Now return ONLY the job description section.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",  # or "gpt-3.5-turbo"
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=1000
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"❌ GPT failed: {e}"

def save_rewritten_resume_docx(rewritten_text, filename):
    doc = Document()
    for line in rewritten_text.split('\n'):
        doc.add_paragraph(line)
    path = os.path.join('uploads', filename)
    doc.save(path)
    return path


def generate_cover_letter(resume_text, job_description):
    prompt = f"""
You are an expert in writing personalized, professional cover letters.

Use the resume and job description below to write a tailored cover letter. Focus on:
- Matching the job’s key responsibilities and qualifications
- Highlighting relevant experience and skills from the resume
- Using a confident and clear tone

Format it with:
- A proper greeting
- 3–4 paragraphs (intro, alignment with role, background, closing)
- A polite, enthusiastic sign-off
- Keep it to one page and no bullet points
- Use a professional tone

--- Resume ---
{resume_text}

--- Job Description ---
{job_description}

Now write the cover letter.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=800
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"❌ GPT failed: {e}"

def save_cover_letter_docx(cover_letter, filename):
    doc = Document()
    for line in cover_letter.split('\n'):
        doc.add_paragraph(line)
    path = os.path.join('uploads', filename)
    doc.save(path)
    return path

@app.route('/download', methods=['POST'])
def download():
    filename = request.form['filename']
    path = os.path.join(UPLOAD_FOLDER, filename)
    @after_this_request
    def cleanup(response):
        try:
            os.remove(path)
        except Exception as e:
            print(f"Cleanup failed: {e}")
        return response

    return send_file(path, as_attachment=True)

@app.route('/upload', methods=['GET', 'POST'])
def upload():
    if request.method == 'GET':
        return render_template('upload.html', year=datetime.now().year)
    
    resume = request.files['resume']
    job_input= request.form['job_description']

    if job_input.startswith("http"):
        try:
            page_text = scrape_webpage(job_input)
            job_description = extract_job_description_from_text(page_text)
            ## Debugging output
            print("🧠 Extracted Job Description from URL:\n", job_description)
        except Exception as e:
            return f"❌ Failed to scrape job description: {str(e)}"
    else:
        job_description = job_input
    if resume:
        file_path = os.path.join(UPLOAD_FOLDER, resume.filename)
        resume.save(file_path)

        # Extract text based on file type
        if resume.filename.endswith('.pdf'):
            resume_text = extract_text_from_pdf(file_path)
        elif resume.filename.endswith('.docx'):
            resume_text = extract_text_from_docx(file_path)
        else:
            return "❌ Unsupported file type."

        rewritten = rewrite_resume(resume_text, job_description)
        rewritten_filename = "rewritten_" + os.path.splitext(resume.filename)[0] + ".docx"
        rewritten_path = save_rewritten_resume_docx(rewritten, rewritten_filename)
        cover_letter = generate_cover_letter(resume_text, job_description)
        cover_letter_filename = "cover_letter_" + os.path.splitext(resume.filename)[0] + ".docx"
        cover_letter_path = save_cover_letter_docx(cover_letter, cover_letter_filename)
        print("📝 COVER LETTER:\n", cover_letter)

        # Schedule file cleanup AFTER sending response
        @after_this_request
        def cleanup(response):
            try:
                os.remove(file_path)
            except Exception as e:
                print(f"Cleanup failed: {e}")
            return response

        return render_template('preview.html', rewritten_text=rewritten, filename=rewritten_filename, cover_letter=cover_letter, cover_letter_filename=cover_letter_filename, year=datetime.now().year)

    return "❌ No file uploaded."

from datetime import datetime

@app.route('/add-job', methods=['GET', 'POST'])
def add_job():
    if request.method == 'POST':
        url = request.form['url']
        status = request.form['status']
        notes = request.form['notes']

        # Defaults
        title = ""
        company = ""
        location = ""
        applied_date = datetime.now().strftime("%Y-%m-%d")

        if url.startswith("http"):
            page_text = scrape_job_posting_text(url)
            extracted = extract_title_and_company_from_text(page_text)

            # Parse the GPT result
            for line in extracted.splitlines():
                if "Job Title:" in line:
                    title = line.split("Job Title:")[-1].strip()
                elif "Company:" in line:
                    company = line.split("Company:")[-1].strip()
                elif "Location:" in line:
                    location = line.split("Location:")[-1].strip()

        job_entry = {
            "url": url,
            "title": title,
            "company": company,
            "location": location,
            "status": status,
            "resume_file": "",
            "cover_letter_file": "",
            "applied_date": applied_date,
            "notes": notes
        }

        try:
            with open('applications.json', 'r') as f:
                applications = json.load(f)
        except (FileNotFoundError, json.decoder.JSONDecodeError):
            applications = []

        applications.append(job_entry)

        with open('applications.json', 'w') as f:
            json.dump(applications, f, indent=4)

        return redirect('/tracker?success=job')

    return render_template('add_job.html', year=datetime.now().year)

@app.route('/tracker')
def tracker():
    try:
        with open('applications.json', 'r') as f:
            applications = json.load(f)
    except (FileNotFoundError, json.decoder.JSONDecodeError):
        applications = []

    return render_template('tracker.html', applications=applications, year=datetime.now().year)

from flask import request, redirect

@app.route('/update-status', methods=['POST'])
def update_status():
    url = request.form['url']
    new_status = request.form['new_status']

    try:
        with open('applications.json', 'r') as f:
            applications = json.load(f)
    except (FileNotFoundError, json.decoder.JSONDecodeError):
        applications = []

    # Update the matching job entry
    for app in applications:
        if app['url'] == url:
            app['status'] = new_status
            break

    # Save updated data
    with open('applications.json', 'w') as f:
        json.dump(applications, f, indent=4)

    return redirect('/tracker?success=status')


def scrape_job_posting_text(url):
    try:
        page = requests.get(url)
        soup = BeautifulSoup(page.text, 'lxml')
        for tag in soup(['script', 'style']):
            tag.decompose()
        return soup.get_text()
    except Exception as e:
        return f"❌ Failed to scrape page: {str(e)}"

def extract_title_and_company_from_text(page_text):
    prompt = f"""
Below is raw text from a job listing webpage. Your task is to extract:

1. Job title (e.g., 'Software Engineer Intern')
2. Company name (e.g., 'Stripe')
3. Location (e.g., 'Remote, USA')
4. The date applied (e.g., '2023-10-01')

Just return the title and company like this:

Job Title: ...
Company: ...

Here is the text:
\"\"\"
{page_text}
\"\"\"
"""

    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=300
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"❌ GPT failed: {e}"

@app.route('/recommend_jobs')
def recommend_jobs():
    # Load tracked applications
    try:
        with open('applications.json', 'r') as f:
            applications = json.load(f)
    except (FileNotFoundError, json.decoder.JSONDecodeError):
        applications = []

    # Extract past job titles as query seeds
    job_titles = list({app['title'] for app in applications if app.get('title')})

    # Use first title as fallback if empty
    if not job_titles:
        job_titles = ["Software Engineer Intern"]

    headers = {
        "X-RapidAPI-Key": os.getenv("JSEARCH_API_KEY"),
        "X-RapidAPI-Host": "jsearch.p.rapidapi.com"
    }

    results = []

    for title in job_titles[:3]:  # Limit to top 3 titles to avoid API spam
        params = {"query": title, "num_pages": 1}
        res = requests.get("https://jsearch.p.rapidapi.com/search", headers=headers, params=params)

        if res.status_code == 200:
            jobs = res.json().get("data", [])
            for job in jobs:
                results.append({
                    "title": job.get("job_title"),
                    "company": job.get("employer_name"),
                    "location": job.get("job_city") or job.get("job_country"),
                    "description": job.get("job_description")[:300] + "...",
                    "url": job.get("job_apply_link")
                })
        else:
            print("❌ Failed JSearch call:", res.text)

    # You'll later sort these based on your matching logic
    return render_template("recommendations.html", jobs=results, year=datetime.now().year)


if __name__ == '__main__':
    app.run(debug=True)